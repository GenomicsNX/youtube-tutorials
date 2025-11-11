# Required libraries installation:
# pip install ollama pypdf2 python-docx

import PyPDF2
from docx import Document
from ollama import Client

OLLAMA_MODEL = "gpt-oss:120b-cloud"
OLLAMA_API_KEY = "***"

def ask_ollama(client, prompt):
    try:
        response = client.chat(
            model=OLLAMA_MODEL,
            messages=[
                {"role": "user", "content": prompt}
            ],
        )
        
        return response.message.content.strip()
    except Exception as e:
        return f"Error: {e}"

def create_client(api_key):
    return Client(
        host="https://ollama.com",
        headers={"Authorization": f"Bearer {api_key}"}
    )


def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def summarize_text(client, text):
    """Send text to the model for summarization"""
    prompt = f"""
        Summarize this text in one paragraph (Persian Language):\n\n{text}
    """
    result = ask_ollama(client, prompt)
    return result

def save_summary_to_word(summary, output_path="summary.docx"):
    """Save summary to Word file"""
    doc = Document()
    doc.add_heading("PDF File Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(output_path)
    print(f"✅ Summary successfully saved to file '{output_path}'!")

if __name__ == "__main__":
    pdf_path = "data.pdf"   # Input PDF file path
    print("📄 Reading PDF file ...")
    text = extract_text_from_pdf(pdf_path)

    print("🧠 Generating summary with Ollama ...")
    client = create_client(OLLAMA_API_KEY)
    summary = summarize_text(client, text)

    print("💾 Saving to Word file ...")
    save_summary_to_word(summary)
